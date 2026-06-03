#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
md2doc.py  v2.0
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Markdown → PDF / Word  本地转换工具
支持：中文 / 数学 LaTeX / 有机化学式（\\ce{}）

PDF  引擎：Playwright Headless Chromium + MathJax（本地优先，CDN 备用）
Word 引擎：python-docx + latex2mathml（纯 lxml，修复 v1 命名空间损坏问题）

PDF 字体主题：
  1. 苹方 + Helvetica Neue（现代）
  2. Noto Serif CJK SC + Times New Roman（学术衬线）
Word 字体：宋体（中文）+ Times New Roman（英文）+ Courier New（代码）
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
依赖安装（一次性）：
  pip install markdown playwright python-docx latex2mathml lxml
  playwright install chromium
  npm install -g mathjax          # 可选，离线 PDF 公式渲染
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
"""

import subprocess, sys, os, re, tempfile, textwrap, zipfile, io

# ─────────────────────────────────────────────────
# § 0  依赖检查
# ─────────────────────────────────────────────────
def _check_deps():
    import importlib.util as _ilu
    missing = [pip for pip, mod in {
        "markdown":     "markdown",
        "playwright":   "playwright",
        "python-docx":  "docx",
        "latex2mathml": "latex2mathml",
        "lxml":         "lxml",
    }.items() if not _ilu.find_spec(mod)]
    if missing:
        print(f"❌ 缺少依赖：{', '.join(missing)}")
        print(f"   请运行：pip install {' '.join(missing)}")
        if "playwright" in missing:
            print("            playwright install chromium")
        sys.exit(1)

_check_deps()

import markdown
import latex2mathml.converter
from lxml import etree as _ET
from playwright.sync_api import sync_playwright
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


# ─────────────────────────────────────────────────
# § 1  剪贴板读取（macOS pbpaste）
# ─────────────────────────────────────────────────
def read_clipboard() -> str:
    print("📋 正在从剪贴板读取内容…")
    text = ""
    if sys.platform == "darwin":
        try:
            r = subprocess.run(["pbpaste"], capture_output=True,
                               text=True, encoding="utf-8")
            text = r.stdout
        except FileNotFoundError:
            print("❌ 未找到 pbpaste，请在 macOS 上运行。")
            sys.exit(1)
    elif sys.platform == "win32":
        try:
            import tkinter as tk
            root = tk.Tk(); root.withdraw()
            text = root.clipboard_get(); root.destroy()
        except Exception as e:
            print(f"❌ 读取剪贴板失败：{e}")
            sys.exit(1)
    else:
        print("❌ 不支持的操作系统，请在 macOS 或 Windows 上运行。")
        sys.exit(1)
    if not text.strip():
        print("❌ 剪贴板为空，请先复制 Markdown 内容。")
        sys.exit(1)
    print(f"✅ 读取成功，共 {len(text.splitlines())} 行。")
    return text


# ─────────────────────────────────────────────────
# § 2  交互提示
# ─────────────────────────────────────────────────
def _ask(prompt: str, valid: dict[str, str], default: str) -> str:
    while True:
        c = input(prompt).strip()
        if c == "": return default
        if c in valid: return valid[c]
        print(f"⚠️  请输入 {' 或 '.join(valid)}。")

def ask_format() -> str:
    print("\n请选择输出格式：\n  1. PDF\n  2. Word (.docx)")
    return _ask("请输入 1 或 2（默认 1）: ", {"1": "pdf", "2": "word"}, "pdf")

def ask_pdf_theme() -> str:
    print("\n请选择 PDF 字体主题：")
    print("  1. 苹方 + Helvetica Neue（现代感）")
    print("  2. Noto Serif CJK SC + Times New Roman（学术衬线）")
    return _ask("请输入 1 或 2（默认 1）: ", {"1": "pingfang", "2": "noto"}, "pingfang")

def ask_filename(ext: str) -> str:
    while True:
        name = input(f"\n请输入输出文件名（不含 {ext}）: ").strip()
        if name:
            return name.removesuffix(ext)
        print("⚠️  文件名不能为空。")

# ── 去除 AI 寒暄开场白 ──
_HR_RE = re.compile(r'^[-*_]{3,}\s*$')                 # 水平分隔线 --- / *** / ___
# 正文结构特征（标题/列表/表格/代码/引用/块级公式），出现即认定不是寒暄
_STRUCT_RE = re.compile(r'^\s*(#{1,6}\s|[-*+]\s|\d+\.\s|\||```|>|\$\$)')
_PREAMBLE_MAX_CHARS = 150                              # 分隔线前内容超过此长度则视为正文

def maybe_strip_ai_preamble(md_text: str) -> str:
    """
    检测并（询问后）去除 AI 在首条水平分隔线前的寒暄开场白。
    判定：首条 ---/***/___ 之前的内容较短、且不含正文结构 → 疑似寒暄。
    """
    lines = md_text.splitlines()
    hr_idx = next((idx for idx, ln in enumerate(lines)
                   if _HR_RE.match(ln.strip())), None)
    if hr_idx is None:
        return md_text                                 # 没有分隔线，不处理

    head = lines[:hr_idx]
    preamble = "\n".join(head).strip()
    if not preamble:
        return md_text                                 # 分隔线前无内容
    if len(preamble) > _PREAMBLE_MAX_CHARS or any(_STRUCT_RE.match(l) for l in head):
        return md_text                                 # 内容多或含正文结构 → 当作正文

    print("\n🤖 检测到首条分隔线前疑似 AI 寒暄开场白：")
    print("   " + "─" * 52)
    for l in preamble.splitlines():
        print(f"   │ {l}")
    print("   " + "─" * 52)
    if _ask("是否删除这段开场白？(y=删除 / n=保留，默认 y): ",
            {"y": "y", "Y": "y", "n": "n", "N": "n"}, "y") == "n":
        print("   ↩️  已保留开场白。")
        return md_text

    rest = lines[hr_idx + 1:]                           # 删掉开场白与该分隔线
    while rest and not rest[0].strip():                 # 顺带去掉其后紧邻空行
        rest.pop(0)
    print("   ✅ 已删除开场白。")
    return "\n".join(rest)


# ─────────────────────────────────────────────────
# § 3  Markdown 解析（保护 LaTeX 避免被破坏）
# ─────────────────────────────────────────────────
# 行内公式：排除货币写法（$120、$360,000）。
# 规则（仿 pandoc）：开 $ 不能前接数字、后接空白；闭 $ 不能前接空白、后接数字。
# 这样 "$120 | $360,000" 不会被误配对成公式，避免吞掉表格 | 列分隔符。
_INLINE_MATH = r'(?<![\d$])\$(?!\s)[^$\n]+?(?<!\s)\$(?!\d)'

def _protect_latex(text: str) -> tuple[str, dict]:
    """
    把 $$...$$  和  $...$  替换为占位符，
    防止 Markdown 解析器将下划线解析为斜体、或把货币 $ 误当公式。
    """
    ph: dict[str, str] = {}
    n = [0]

    def _store(m: re.Match) -> str:
        key = f"\x00L{n[0]}\x00"; n[0] += 1
        ph[key] = m.group(0); return key

    text = re.sub(r'\$\$[\s\S]+?\$\$', _store, text)   # 块级优先
    text = re.sub(_INLINE_MATH,        _store, text)   # 再行内
    return text, ph

def md_to_html_body(md_text: str) -> str:
    protected, ph = _protect_latex(md_text)
    exts = ["extra", "codehilite", "toc", "nl2br", "sane_lists"]
    cfg  = {"codehilite": {"guess_lang": False, "noclasses": True}}
    html = markdown.markdown(protected, extensions=exts, extension_configs=cfg)
    for k, v in ph.items():
        html = html.replace(k, v)
    return html


# ─────────────────────────────────────────────────
# § 4  PDF 生成
# ─────────────────────────────────────────────────

# ── 4a. MathJax 本地路径检测 ──
def _find_local_mathjax() -> str | None:
    candidates = [
        os.path.expanduser("~/node_modules/mathjax/es5/tex-chtml.js"),
        os.path.join(os.getcwd(), "node_modules/mathjax/es5/tex-chtml.js"),
        "/usr/local/lib/node_modules/mathjax/es5/tex-chtml.js",
        "/opt/homebrew/lib/node_modules/mathjax/es5/tex-chtml.js",
        os.path.expandvars(r"%APPDATA%\npm\node_modules\mathjax\es5\tex-chtml.js"),
        os.path.expandvars(r"%ProgramFiles%\nodejs\node_modules\mathjax\es5\tex-chtml.js"),
    ]
    return next((p for p in candidates if os.path.exists(p)), None)

# ── 4b. CSS 字体主题 ──
_THEMES: dict[str, dict] = {
    "pingfang": dict(
        body='"PingFang SC","Heiti SC","Microsoft YaHei","微软雅黑","Helvetica Neue",Arial,sans-serif',
        head='"PingFang SC","Heiti SC","Microsoft YaHei","微软雅黑","Helvetica Neue",Arial,sans-serif',
        code='"Monaco","Menlo","Consolas","Courier New",monospace',
        h1c="#2c3e50", h2c="#bdc3c7", lnk="#2980b9",
        qbg="#eaf4fb", qbd="#3498db", cbd="#3498db", thbg="#2c3e50",
    ),
    "noto": dict(
        body='"Noto Serif CJK SC","Noto Serif SC","SimSun","宋体","Times New Roman",serif',
        head='"Noto Serif CJK SC","Noto Serif SC","SimSun","宋体","Times New Roman",serif',
        code='"Noto Sans Mono CJK SC","Consolas","Menlo","Courier New",monospace',
        h1c="#1a3a5c", h2c="#95a5a6", lnk="#1a5276",
        qbg="#f4f6f7", qbd="#1a5276", cbd="#1a5276", thbg="#1a3a5c",
    ),
}

def _build_css(t: dict) -> str:
    return f"""
*,*::before,*::after{{box-sizing:border-box;margin:0;padding:0}}
body{{
  font-family:{t['body']};font-size:10.5pt;line-height:1.8;
  color:#1a1a1a;background:#fff;
  width:170mm;margin:0 auto;padding:20mm 0 25mm;
  word-wrap:break-word;overflow-wrap:break-word;
}}
h1,h2,h3,h4,h5,h6{{
  font-family:{t['head']};font-weight:700;
  margin-top:1.6em;margin-bottom:.5em;line-height:1.4;
  page-break-after:avoid;
}}
h1{{font-size:20pt;border-bottom:2px solid {t['h1c']};padding-bottom:6px}}
h2{{font-size:15pt;border-bottom:1px solid {t['h2c']};padding-bottom:4px}}
h3{{font-size:13pt}} h4{{font-size:11.5pt}}
h5,h6{{font-size:10.5pt;color:#555}}
p{{margin:.6em 0 .8em}}
strong{{font-weight:700}} em{{font-style:italic}}
a{{color:{t['lnk']};text-decoration:none}} a:hover{{text-decoration:underline}}
code{{
  font-family:{t['code']};font-size:9pt;
  background:#f4f4f4;border:1px solid #ddd;border-radius:3px;padding:1px 5px;
}}
pre{{
  font-family:{t['code']};font-size:8.5pt;
  background:#f8f8f8;border:1px solid #ddd;
  border-left:4px solid {t['cbd']};border-radius:4px;
  padding:12px 16px;margin:1em 0;
  white-space:pre-wrap;word-break:break-all;page-break-inside:avoid;
}}
pre code{{background:none;border:none;padding:0;font-size:inherit}}
blockquote{{
  border-left:4px solid {t['qbd']};background:{t['qbg']};
  margin:1em 0;padding:10px 16px;border-radius:0 4px 4px 0;
  page-break-inside:avoid;
}}
blockquote p{{margin:0}}
ul,ol{{margin:.5em 0 .8em 1.6em;padding:0}}
li{{margin:.3em 0}} li>ul,li>ol{{margin:.2em 0 .2em 1.2em}}
/* 三线表：仅顶/底两条粗线 + 表头下一条细线，无竖线、无斑马纹 */
table{{
  width:100%;border-collapse:collapse;table-layout:fixed;
  margin:1.2em 0;font-size:9.5pt;
  border-top:1.5pt solid #000;border-bottom:1.5pt solid #000;
}}
thead{{display:table-header-group}}      /* 跨页时表头自动重复 */
thead tr{{background:transparent;color:inherit;border-bottom:1pt solid #000}}
th,td{{
  border:none;padding:6px 9px;text-align:left;vertical-align:top;
  /* 关键：允许长内容换行，防止超宽被裁剪导致数据丢失 */
  word-break:break-word;overflow-wrap:anywhere;white-space:normal;
}}
th{{font-weight:700}}
tr{{page-break-inside:avoid}}
hr{{border:none;border-top:2px solid #ecf0f1;margin:2em 0}}
img{{max-width:100%;height:auto;display:block;margin:1em auto;border-radius:4px}}
mjx-container{{overflow-x:auto}}
.MathJax_Display,mjx-container[display="true"]{{
  margin:1em 0!important;page-break-inside:avoid;
}}
@media print{{
  body{{width:100%;padding:0}}
  pre,blockquote,figure{{page-break-inside:avoid}}
  h1,h2,h3,h4{{page-break-after:avoid}}
  p,li{{orphans:3;widows:3}}
}}"""

# ── 4c. HTML 组装 ──
_MJ_CFG = r"""window.MathJax={
  tex:{
    inlineMath:[['$','$'],['\\(','\\)']],
    displayMath:[['$$','$$'],['\\[','\\]']],
    processEscapes:true,processEnvironments:true,
    tags:'ams',packages:{'[+]':['mhchem']}
  },
  options:{skipHtmlTags:['script','noscript','style','textarea','pre']},
  loader:{load:['[tex]/mhchem']}
};"""

def build_full_html(body_html: str, theme: str = "pingfang") -> str:
    css = _build_css(_THEMES[theme])
    local = _find_local_mathjax()
    if local:
        print(f"  ℹ️  使用本地 MathJax：{local}")
        mj_src = f"file://{local}"
    else:
        print("  ℹ️  未找到本地 MathJax，使用 CDN（需联网）。")
        print("     离线安装：npm install -g mathjax")
        mj_src = "https://cdn.jsdelivr.net/npm/mathjax@3/es5/tex-chtml.js"
    return textwrap.dedent(f"""\
        <!DOCTYPE html><html lang="zh-CN">
        <head>
          <meta charset="UTF-8"><title>Export</title>
          <style>{css}</style>
          <script>{_MJ_CFG}</script>
          <script async src="{mj_src}"></script>
        </head>
        <body>{body_html}</body>
        </html>""")

# ── 4d. Playwright 渲染 ──
def html_to_pdf(html: str, out: str) -> None:
    print("🚀 启动无头浏览器…")
    with tempfile.NamedTemporaryFile(
            mode="w", suffix=".html", delete=False, encoding="utf-8") as f:
        f.write(html); tmp = f.name
    try:
        with sync_playwright() as pw:
            browser = pw.chromium.launch(headless=True)
            page    = browser.new_page()
            page.goto(f"file://{tmp}", wait_until="networkidle")
            print("⏳ 等待 MathJax 渲染（最长 30 秒）…")
            try:
                page.wait_for_function(
                    "()=>window.MathJax&&window.MathJax.startup"
                    "&&typeof window.MathJax.startup.promise==='object'",
                    timeout=30_000)
                page.evaluate("()=>window.MathJax.startup.promise")
                page.wait_for_timeout(1_000)
            except Exception:
                print("⚠️  MathJax 超时，继续生成（公式可能显示为源码）。")
            print("🖨️  生成 PDF…")
            page.pdf(
                path=out, format="A4", print_background=True,
                margin={"top":"20mm","bottom":"25mm",
                        "left":"20mm","right":"20mm"})
            browser.close()
    finally:
        os.unlink(tmp)


# ─────────────────────────────────────────────────
# § 5  Word 生成（纯 lxml，避免命名空间损坏）
# ─────────────────────────────────────────────────
_W_ZH    = "宋体"
_W_EN    = "Times New Roman"
_W_CODE  = "Courier New"

# ── 5a. 字体工具 ──
def _set_run_fonts(run, zh: str = _W_ZH, en: str = _W_EN) -> None:
    """用纯 lxml 同时设置东西文字体（不使用 OxmlElement 避免混用）。"""
    run.font.name = en
    rPr = run._r.get_or_add_rPr()
    for old in rPr.findall(qn("w:rFonts")):
        rPr.remove(old)
    rF = _ET.SubElement(rPr, qn("w:rFonts"))
    rF.set(qn("w:ascii"),    en)
    rF.set(qn("w:hAnsi"),    en)
    rF.set(qn("w:eastAsia"), zh)
    rF.set(qn("w:cs"),       zh)
    rPr.insert(0, rF)

def _set_doc_defaults(doc: Document) -> None:
    """文档默认字体（宋体 + Times New Roman）与 1.5 倍行距。"""
    normal = doc.styles["Normal"]
    normal.font.name = _W_EN
    normal.font.size = Pt(11)
    rPr = normal.element.get_or_add_rPr()
    for old in rPr.findall(qn("w:rFonts")):
        rPr.remove(old)
    rF = _ET.SubElement(rPr, qn("w:rFonts"))
    rF.set(qn("w:ascii"),    _W_EN); rF.set(qn("w:hAnsi"),    _W_EN)
    rF.set(qn("w:eastAsia"), _W_ZH); rF.set(qn("w:cs"),       _W_ZH)
    rPr.insert(0, rF)
    pPr = normal.element.get_or_add_pPr()
    for old in pPr.findall(qn("w:spacing")):
        pPr.remove(old)
    spc = _ET.SubElement(pPr, qn("w:spacing"))
    spc.set(qn("w:line"),     "360")   # 240 = 单倍，360 = 1.5 倍
    spc.set(qn("w:lineRule"), "auto")

# ── 5b. LaTeX → OMML / MathML（纯 lxml）──
def _try_office_xslt(ml_str: str) -> "_ET._Element | None":
    """
    利用 Office for Mac 自带 MML2OMML.XSL 将 MathML 转为原生 OMML。
    原生 OMML 在 Word 中完全可编辑，兼容性最佳。
    """
    xsl_candidates = [
        "/Applications/Microsoft Word.app/Contents/Resources/MML2OMML.XSL",
        os.path.expanduser(
            "~/Library/Application Support/Microsoft/Office/MML2OMML.XSL"),
        r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL",
        r"C:\Program Files (x86)\Microsoft Office\root\Office16\MML2OMML.XSL",
        r"C:\Program Files\Microsoft Office\root\Office15\MML2OMML.XSL",
        r"C:\Program Files (x86)\Microsoft Office\root\Office15\MML2OMML.XSL",
    ]
    xsl = next((p for p in xsl_candidates if os.path.exists(p)), None)
    if xsl is None:
        return None
    try:
        xslt      = _ET.XSLT(_ET.parse(xsl))
        ml_doc    = _ET.fromstring(ml_str.encode())
        omml_tree = xslt(ml_doc)
        return omml_tree.getroot()
    except Exception:
        return None

# ── 内置 OMML 生成器（直接 LaTeX → OMML，无需 Office XSL）──
_M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"

_L_SYM: dict[str, str] = {
    # 希腊字母小写
    'alpha':'α','beta':'β','gamma':'γ','delta':'δ','epsilon':'ε',
    'varepsilon':'ε','zeta':'ζ','eta':'η','theta':'θ','vartheta':'ϑ',
    'iota':'ι','kappa':'κ','lambda':'λ','mu':'μ','nu':'ν','xi':'ξ',
    'pi':'π','varpi':'ϖ','rho':'ρ','varrho':'ϱ','sigma':'σ',
    'varsigma':'ς','tau':'τ','upsilon':'υ','phi':'φ','varphi':'φ',
    'chi':'χ','psi':'ψ','omega':'ω',
    # 希腊字母大写
    'Gamma':'Γ','Delta':'Δ','Theta':'Θ','Lambda':'Λ','Xi':'Ξ',
    'Pi':'Π','Sigma':'Σ','Upsilon':'Υ','Phi':'Φ','Psi':'Ψ','Omega':'Ω',
    # 运算符与关系
    'times':'×','div':'÷','pm':'±','mp':'∓','cdot':'·','ast':'∗',
    'leq':'≤','geq':'≥','neq':'≠','ne':'≠','approx':'≈','equiv':'≡',
    'sim':'∼','cong':'≅','propto':'∝','ll':'≪','gg':'≫',
    # 箭头
    'rightarrow':'→','leftarrow':'←','Rightarrow':'⇒','Leftarrow':'⇐',
    'leftrightarrow':'↔','Leftrightarrow':'⇔','to':'→',
    'uparrow':'↑','downarrow':'↓',
    # 杂项
    'infty':'∞','partial':'∂','nabla':'∇','forall':'∀','exists':'∃',
    'in':'∈','notin':'∉','subset':'⊂','supset':'⊃',
    'subseteq':'⊆','supseteq':'⊇','cup':'∪','cap':'∩',
    'emptyset':'∅','varnothing':'∅',
    'cdots':'⋯','ldots':'…','vdots':'⋮','ddots':'⋱',
    'hbar':'ℏ','ell':'ℓ','Re':'ℜ','Im':'ℑ','aleph':'ℵ',
    'oplus':'⊕','otimes':'⊗','bullet':'•','circ':'∘',
    # 函数名（直立体）
    'sin':'sin','cos':'cos','tan':'tan','cot':'cot','sec':'sec','csc':'csc',
    'arcsin':'arcsin','arccos':'arccos','arctan':'arctan',
    'sinh':'sinh','cosh':'cosh','tanh':'tanh',
    'log':'log','ln':'ln','exp':'exp','lim':'lim',
    'max':'max','min':'min','sup':'sup','inf':'inf',
    'det':'det','deg':'deg','dim':'dim','ker':'ker','gcd':'gcd',
    'arg':'arg','mod':'mod','Pr':'Pr',
    # 空白
    'quad':' ','qquad':'  ',',':' ',';':' ','!':'',
}
_FUNC_SET = {'sin','cos','tan','cot','sec','csc','arcsin','arccos','arctan',
             'sinh','cosh','tanh','log','ln','exp','lim','max','min','sup',
             'inf','det','deg','dim','ker','gcd','arg','mod','Pr'}
_NARY_MAP: dict[str, tuple[str, bool]] = {
    'sum':('∑',True),'prod':('∏',True),'coprod':('∐',True),
    'int':('∫',False),'iint':('∬',False),'iiint':('∭',False),
    'oint':('∮',False),'bigcup':('⋃',True),'bigcap':('⋂',True),
    'bigvee':('⋁',True),'bigwedge':('⋀',True),
}

def _mel(tag: str) -> _ET._Element:
    return _ET.Element(f"{{{_M_NS}}}{tag}")

def _mval(el: _ET._Element, attr: str, val: str) -> None:
    el.set(f"{{{_M_NS}}}{attr}", val)

def _mrun(text: str, italic: bool | None = None) -> _ET._Element:
    if italic is None:
        italic = len(text) == 1 and text.isalpha() and text.isascii()
    r = _mel("r")
    rPr = _mel("rPr"); sty = _mel("sty"); _mval(sty, "val", "i" if italic else "p")
    rPr.append(sty); r.append(rPr)
    t = _mel("t"); t.text = text; r.append(t)
    return r

def _mext(parent: _ET._Element, children) -> None:
    for c in (children or []):
        if isinstance(c, list): _mext(parent, c)
        elif c is not None:     parent.append(c)

def _tok(s: str) -> list[tuple[str, str]]:
    toks, i = [], 0
    while i < len(s):
        c = s[i]
        if c == '\\' and i + 1 < len(s):
            nc = s[i + 1]
            if nc.isalpha():
                j = i + 2
                while j < len(s) and s[j].isalpha(): j += 1
                toks.append(('cmd', s[i+1:j])); i = j
                while i < len(s) and s[i] == ' ': i += 1
            else:
                toks.append(('cmd', nc)); i += 2
        elif c in '{}^_&':
            toks.append(('sp', c)); i += 1
        elif c != ' ':
            toks.append(('ch', c)); i += 1
        else:
            i += 1
    return toks

class _MP:
    """极简 LaTeX → OMML 递归下降解析器。"""
    def __init__(self, t: list): self._t = t; self._i = 0
    def _p(self): return self._t[self._i] if self._i < len(self._t) else None
    def _n(self): r = self._t[self._i]; self._i += 1; return r

    def seq(self) -> list:
        out = []
        while self._i < len(self._t):
            if self._p() in (('sp', '}'), None): break
            base = self._atom()
            if base is None: continue
            bl = base if isinstance(base, list) else [base]
            sub = sup = None
            while self._i < len(self._t):
                np = self._p()
                if np == ('sp', '^') and sup is None: self._n(); sup = self._ga()
                elif np == ('sp', '_') and sub is None: self._n(); sub = self._ga()
                else: break
            if sub is None and sup is None:
                out.extend(bl)
            elif sub is not None and sup is not None:
                el = _mel("sSubSup"); e = _mel("e"); _mext(e, bl)
                s = _mel("sub"); _mext(s, sub); sp2 = _mel("sup"); _mext(sp2, sup)
                el.extend([e, s, sp2]); out.append(el)
            elif sub is not None:
                el = _mel("sSub"); e = _mel("e"); _mext(e, bl)
                s = _mel("sub"); _mext(s, sub); el.extend([e, s]); out.append(el)
            else:
                el = _mel("sSup"); e = _mel("e"); _mext(e, bl)
                s = _mel("sup"); _mext(s, sup); el.extend([e, s]); out.append(el)
        return out

    def _ga(self) -> list:
        if self._p() == ('sp', '{'): return self._grp()
        a = self._atom()
        return [a] if a is not None else []

    def _grp(self) -> list:
        self._n()
        c = self.seq()
        if self._p() == ('sp', '}'): self._n()
        return c

    def _atom(self):
        t = self._p()
        if t is None: return None
        typ, val = t
        if typ == 'sp':
            if val == '{':
                c = self._grp(); return c if len(c) != 1 else c[0]
            self._n(); return None
        self._n()
        if typ == 'ch': return _mrun(val)
        return self._cmd(val)

    def _cmd(self, cmd: str):
        if cmd in _NARY_MAP: return self._nary(*_NARY_MAP[cmd])
        if cmd == 'frac':
            n, d = self._grp(), self._grp()
            f = _mel("f"); num = _mel("num"); _mext(num, n)
            den = _mel("den"); _mext(den, d); f.extend([num, den]); return f
        if cmd == 'sqrt':
            deg = None
            if self._p() == ('ch', '['):
                self._n(); dc = []
                while self._p() and self._p() != ('ch', ']'): dc.append(self._n()[1])
                if self._p(): self._n()
                deg = ''.join(dc)
            c = self._grp(); rad = _mel("rad"); rpr = _mel("radPr")
            if not deg: dh = _mel("degHide"); _mval(dh, "val", "1"); rpr.append(dh)
            rad.append(rpr); dg = _mel("deg")
            if deg: dg.append(_mrun(deg, False))
            rad.append(dg); e = _mel("e"); _mext(e, c); rad.append(e); return rad
        if cmd == 'left':  return self._delim()
        if cmd == 'right': return None
        # 直立文本
        if cmd in ('text', 'mathrm', 'textrm', 'mbox', 'mathop', 'operatorname',
                   'mathnormal'):
            c = self._grp()
            txt = ''.join(x.findtext(f"{{{_M_NS}}}t") or '' for x in (c or [])
                          if hasattr(x, 'tag') and x.tag == f"{{{_M_NS}}}r")
            return _mrun(txt or '', False)
        # 样式命令：消化参数，内容原样返回
        if cmd in ('mathbf', 'boldsymbol', 'bm', 'mathbb', 'mathcal',
                   'mathscr', 'mathfrak', 'mathit', 'mathsf', 'mathtt',
                   'mathring', 'acute', 'grave', 'breve', 'check',
                   'mathlarger', 'mathsmaller', 'textbf', 'textit'):
            return self._grp()
        if cmd in ('overline', 'bar'):       return self._acc('̄')
        if cmd in ('hat', 'widehat'):        return self._acc('̂')
        if cmd in ('tilde', 'widetilde'):    return self._acc('̃')
        if cmd in ('vec', 'overrightarrow'): return self._acc('⃗')
        if cmd in ('dot', 'ddot'):
            c = self._grp(); acc = _mel("acc"); pr = _mel("accPr")
            ch = _mel("chr"); _mval(ch, "val", '̈' if cmd == 'ddot' else '̇')
            pr.append(ch); acc.append(pr)
            e = _mel("e"); _mext(e, c); acc.append(e); return acc
        if cmd == 'underbrace':
            c = self._grp(); lim = _mel("limLow"); e = _mel("e"); _mext(e, c)
            lim.append(e); lim.append(_mel("lim")); return lim
        if cmd == 'overbrace':
            c = self._grp(); lim = _mel("limUpp"); e = _mel("e"); _mext(e, c)
            lim.append(e); lim.append(_mel("lim")); return lim
        if cmd in ('stackrel', 'overset'):
            top = self._grp(); base = self._grp()
            lim = _mel("limUpp"); e = _mel("e"); _mext(e, base)
            l = _mel("lim"); _mext(l, top)
            lim.append(e); lim.append(l); return lim
        if cmd == 'underset':
            bot = self._grp(); base = self._grp()
            lim = _mel("limLow"); e = _mel("e"); _mext(e, base)
            l = _mel("lim"); _mext(l, bot)
            lim.append(e); lim.append(l); return lim
        if cmd in ('cfrac', 'dfrac', 'tfrac'):
            n, d = self._grp(), self._grp()
            f = _mel("f"); num = _mel("num"); _mext(num, n)
            den = _mel("den"); _mext(den, d); f.extend([num, den]); return f
        if cmd == 'binom':
            n, d = self._grp(), self._grp()
            f = _mel("f"); fPr = _mel("fPr"); _mval(fPr, "type", "noBar")
            f.append(fPr); num = _mel("num"); _mext(num, n)
            den = _mel("den"); _mext(den, d); f.extend([num, den])
            d_el = _mel("d"); dPr = _mel("dPr")
            bc = _mel("begChr"); _mval(bc, "val", "("); dPr.append(bc)
            ec = _mel("endChr"); _mval(ec, "val", ")"); dPr.append(ec)
            d_el.append(dPr); e = _mel("e"); e.append(f); d_el.append(e); return d_el
        if cmd == 'not':
            base = self._ga()
            acc = _mel("acc"); pr = _mel("accPr")
            ch = _mel("chr"); _mval(ch, "val", "̸"); pr.append(ch); acc.append(pr)
            e = _mel("e"); _mext(e, base); acc.append(e); return acc
        if cmd == 'begin': return self._mat(self._gtxt())
        if cmd == 'end':   self._gtxt(); return None
        # 符号表查找
        if cmd in _L_SYM:
            sym = _L_SYM[cmd]
            return _mrun(sym, cmd not in _FUNC_SET) if sym else None
        # 无参数的排版控制指令：静默忽略
        if cmd in ('\\', 'displaystyle', 'textstyle', 'scriptstyle',
                   'scriptscriptstyle', 'limits', 'nolimits', 'notag',
                   'nonumber', 'allowbreak', 'thinspace', 'medspace',
                   'thickspace', 'negthinspace', 'negmedspace', 'negthickspace',
                   'hfil', 'hfill', 'noindent', 'cr', 'relax',
                   'big', 'Big', 'bigg', 'Bigg',
                   'bigl', 'Bigl', 'biggl', 'Biggl',
                   'bigr', 'Bigr', 'biggr', 'Biggr',
                   'middle', 'over', 'choose'):
            return None
        # 吃掉一个参数并忽略的命令
        if cmd in ('color', 'textcolor', 'hspace', 'vspace', 'label',
                   'tag', 'phantom', 'hphantom', 'vphantom', 'smash',
                   'raisebox', 'rule', 'kern', 'mkern', 'mskip'):
            if self._p() == ('sp', '{'): self._grp()
            # textcolor 还有第二个参数（文字内容），返回它
            if cmd == 'textcolor' and self._p() == ('sp', '{'):
                return self._grp()
            return None
        # 未知命令：若后跟 {…} 则返回其内容，否则静默跳过（不输出反斜杠）
        if self._p() == ('sp', '{'):
            return self._grp()
        return None

    def _acc(self, ch_val: str):
        c = self._grp(); acc = _mel("acc"); pr = _mel("accPr")
        ch = _mel("chr"); _mval(ch, "val", ch_val); pr.append(ch); acc.append(pr)
        e = _mel("e"); _mext(e, c); acc.append(e); return acc

    def _nary(self, char: str, has_lim: bool):
        sub = sup = None
        while self._i < len(self._t):
            np = self._p()
            if np == ('sp', '_') and sub is None: self._n(); sub = self._ga()
            elif np == ('sp', '^') and sup is None: self._n(); sup = self._ga()
            else: break
        ny = _mel("nary"); pr = _mel("naryPr")
        ch = _mel("chr"); _mval(ch, "val", char); pr.append(ch)
        if not has_lim: ll = _mel("limLoc"); _mval(ll, "val", "subSup"); pr.append(ll)
        ny.append(pr)
        s = _mel("sub");  _mext(s, sub or []);  ny.append(s)
        sp2 = _mel("sup"); _mext(sp2, sup or []); ny.append(sp2)
        b = self._ga(); e = _mel("e"); _mext(e, b); ny.append(e); return ny

    def _delim(self):
        oc = self._dc(); cont = []
        while self._i < len(self._t):
            if self._p() == ('cmd', 'right'): self._n(); break
            a = self._atom()
            if a is not None: (cont.extend(a) if isinstance(a, list) else cont.append(a))
        cc = self._dc()
        d = _mel("d"); pr = _mel("dPr")
        bc = _mel("begChr"); _mval(bc, "val", oc); pr.append(bc)
        ec = _mel("endChr"); _mval(ec, "val", cc); pr.append(ec)
        d.append(pr); e = _mel("e"); _mext(e, cont); d.append(e); return d

    def _dc(self) -> str:
        t = self._p()
        if t is None: return ''
        typ, val = t; self._n()
        if typ == 'ch':
            return {'(':'(', ')':')', '[':'[', ']':']', '|':'|', '.':''}.get(val) or val
        return {'langle':'⟨','rangle':'⟩','|':'‖','lbrace':'{','rbrace':'}',
                'lfloor':'⌊','rfloor':'⌋','lceil':'⌈','rceil':'⌉'}.get(val, '')

    def _gtxt(self) -> str:
        if self._p() == ('sp', '{'): self._n()
        ch = []
        while self._p() and self._p() != ('sp', '}'): ch.append(self._n()[1])
        if self._p() == ('sp', '}'): self._n()
        return ''.join(ch)

    def _mat(self, env: str):
        dm = {'pmatrix':('(',')'), 'bmatrix':('[',']'), 'vmatrix':('|','|'),
              'Vmatrix':('‖','‖'), 'cases':('{',''), 'matrix':('','')}
        oc, cc = dm.get(env, ('(', ')'))
        rows, row, cell = [], [], []
        while self._i < len(self._t):
            t = self._p()
            if t == ('cmd', 'end'): self._n(); self._gtxt(); break
            if t == ('sp', '&'):   self._n(); row.append(cell); cell = []; continue
            if t == ('cmd', '\\'): self._n(); row.append(cell); rows.append(row); row = []; cell = []; continue
            a = self._atom()
            if a is not None: (cell.extend(a) if isinstance(a, list) else cell.append(a))
        if cell: row.append(cell)
        if row:  rows.append(row)
        m = _mel("m"); m.append(_mel("mPr"))
        for r in rows:
            mr = _mel("mr")
            for cl in r:
                e = _mel("e"); _mext(e, cl); mr.append(e)
            m.append(mr)
        if not oc: return m
        d = _mel("d"); pr = _mel("dPr")
        bc = _mel("begChr"); _mval(bc, "val", oc); pr.append(bc)
        ec = _mel("endChr"); _mval(ec, "val", cc); pr.append(ec)
        d.append(pr); e = _mel("e"); e.append(m); d.append(e); return d


def _build_omml(latex: str) -> "_ET._Element | None":
    """LaTeX → m:oMath，纯 Python 实现，无需 Office XSL。"""
    expr = re.sub(r'^\$+|\$+$', '', latex).strip()
    if not expr:
        return None
    try:
        kids = _MP(_tok(expr)).seq()
        o = _mel("oMath"); _mext(o, kids); return o
    except Exception:
        return None


def _latex_to_omath(latex: str) -> "_ET._Element | None":
    """LaTeX → m:oMath。优先 Office XSL，备用内置生成器。"""
    expr = re.sub(r'^\$+|\$+$', '', latex).strip()
    if not expr:
        return None
    # 优先：Office XSL（最佳兼容性）
    try:
        ml_str = latex2mathml.converter.convert(expr)
        omml = _try_office_xslt(ml_str)
        if omml is not None:
            return omml
    except Exception:
        pass
    # 备用：内置生成器
    return _build_omml(latex)

def _append_formula(para, latex: str) -> None:
    """将公式插入段落；所有路径均失败时退化为等宽纯文本（去掉 $）。"""
    el = _latex_to_omath(latex)
    if el is not None:
        para._p.append(el)
    else:
        expr = re.sub(r'^\$+|\$+$', '', latex).strip()
        r = para.add_run(expr or latex)
        r.font.name = _W_CODE
        r.font.size = Pt(9.5)

# ── 5c. 行内混合文本解析 ──
_INLINE_RE = re.compile(
    r'(`[^`]+`)'                                      # `code`
    r'|((?<![\d$])\$(?!\s)[^$\n]+?(?<!\s)\$(?!\d))'   # $LaTeX$（排除货币 $）
    r'|(\*\*[^*]+\*\*)'                               # **bold**
    r'|(\*[^*]+\*)'                                   # *italic*
    r'|([^`$*\n]+|\$)'                                # 普通文本（含残留货币 $）
)

def _add_inline(para, text: str) -> None:
    for m in _INLINE_RE.finditer(text):
        c, lat, bo, it, pl = m.groups()
        if c:
            r = para.add_run(c[1:-1])
            r.font.name = _W_CODE; r.font.size = Pt(9.5)
        elif lat:
            _append_formula(para, lat)
        elif bo:
            r = para.add_run(bo[2:-2]); r.bold   = True; _set_run_fonts(r)
        elif it:
            r = para.add_run(it[1:-1]); r.italic = True; _set_run_fonts(r)
        elif pl:
            r = para.add_run(pl); _set_run_fonts(r)

# ── 5d. 代码块 ──
def _add_code_block(doc: Document, code_text: str) -> None:
    para = doc.add_paragraph()
    pPr  = para._p.get_or_add_pPr()
    # 蓝色左边框（w:pBdr 必须在 w:shd 之前，符合 OOXML 子元素顺序）
    pBdr = _ET.SubElement(pPr, qn("w:pBdr"))
    lft  = _ET.SubElement(pBdr, qn("w:left"))
    lft.set(qn("w:val"), "single"); lft.set(qn("w:sz"),    "12")
    lft.set(qn("w:space"), "12");   lft.set(qn("w:color"), "2F7BBF")
    # 灰色背景（w:shd 在 w:pBdr 之后）
    shd = _ET.SubElement(pPr, qn("w:shd"))
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4F4F4")
    # 段落间距
    spc = _ET.SubElement(pPr, qn("w:spacing"))
    spc.set(qn("w:before"), "100"); spc.set(qn("w:after"), "100")
    # run
    run = para.add_run(code_text)
    run.font.size = Pt(9)
    rPr = run._r.get_or_add_rPr()
    # 先移除 python-docx 可能已插入的 w:rFonts，避免重复子元素
    for old in rPr.findall(qn("w:rFonts")):
        rPr.remove(old)
    rF  = _ET.SubElement(rPr, qn("w:rFonts"))
    rF.set(qn("w:ascii"),    _W_CODE)
    rF.set(qn("w:hAnsi"),    _W_CODE)
    rF.set(qn("w:eastAsia"), _W_CODE)
    rPr.insert(0, rF)

# ── 5f. 表格 ──────────────────────────────────────────────
def _parse_table_row(line: str) -> list[str]:
    s = line.strip()
    if s.startswith('|'): s = s[1:]
    if s.endswith('|'):   s = s[:-1]
    return [c.strip() for c in s.split('|')]

def _is_sep_row(cells: list[str]) -> bool:
    return bool(cells) and all(re.match(r'^:?-+:?$', c) for c in cells if c)

def _col_align(sep: str) -> WD_ALIGN_PARAGRAPH:
    s = sep.strip()
    if s.startswith(':') and s.endswith(':'): return WD_ALIGN_PARAGRAPH.CENTER
    if s.endswith(':'):                        return WD_ALIGN_PARAGRAPH.RIGHT
    return WD_ALIGN_PARAGRAPH.LEFT

def _set_table_three_line(tbl) -> None:
    """三线表：表格仅保留顶/底两条粗线，去除竖线与内部横线。"""
    tblPr = tbl._tbl.tblPr
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    borders = _ET.SubElement(tblPr, qn("w:tblBorders"))
    def _bd(name: str, sz: int, val: str = "single"):
        e = _ET.SubElement(borders, qn(f"w:{name}"))
        e.set(qn("w:val"), val);   e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), "000000")
    _bd("top", 12); _bd("bottom", 12)                       # 1.5pt 粗线
    _bd("left", 0, "nil"); _bd("right", 0, "nil")
    _bd("insideH", 0, "nil"); _bd("insideV", 0, "nil")
    # 表格宽度占满版心（100%），并允许按内容自适应，避免超宽裁剪
    for old in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old)
    tblW = _ET.SubElement(tblPr, qn("w:tblW"))
    tblW.set(qn("w:w"), "5000"); tblW.set(qn("w:type"), "pct")

def _set_row_props(row, *, header_repeat: bool = False,
                   cant_split: bool = True) -> None:
    """优化表格跨页显示：
    - cant_split：禁止单行内容被从中间拆到两页；
    - header_repeat：表头行在每个分页页面顶部自动重复。
    OOXML trPr 子元素顺序：cantSplit 在 tblHeader 之前。
    """
    trPr = row._tr.get_or_add_trPr()
    for tag in (qn("w:cantSplit"), qn("w:tblHeader")):
        for old in trPr.findall(tag):
            trPr.remove(old)
    if cant_split:
        _ET.SubElement(trPr, qn("w:cantSplit")).set(qn("w:val"), "true")
    if header_repeat:
        _ET.SubElement(trPr, qn("w:tblHeader")).set(qn("w:val"), "true")

def _set_cell_header_border(cell) -> None:
    """表头单元格底部细线（1pt），构成三线表中间那条线。"""
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(old)
    # OOXML 子元素顺序：tcBorders 必须在 shd / tcMar 之前
    tb = _ET.Element(qn("w:tcBorders"))
    bot = _ET.SubElement(tb, qn("w:bottom"))
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "8")
    bot.set(qn("w:space"), "0");    bot.set(qn("w:color"), "000000")
    tcPr.insert(0, tb)

def _add_table_block(doc: Document,
                     header: list[str] | None,
                     body: list[list[str]],
                     aligns: list | None) -> None:
    all_rows = ([header] if header else []) + body
    if not all_rows:
        return
    num_cols = max(len(r) for r in all_rows)
    tbl = doc.add_table(rows=len(all_rows), cols=num_cols)
    tbl.autofit = True          # 按内容/版心自适应列宽，长内容自动换行，防止数据丢失
    tbl.allow_autofit = True
    _set_table_three_line(tbl)

    for ri, row_data in enumerate(all_rows):
        is_hdr = (header is not None and ri == 0)
        _set_row_props(tbl.rows[ri], header_repeat=is_hdr, cant_split=True)
        for ci in range(num_cols):
            cell = tbl.rows[ri].cells[ci]
            para = cell.paragraphs[0]
            if aligns and ci < len(aligns):
                para.alignment = aligns[ci]
            txt = row_data[ci] if ci < len(row_data) else ""
            _add_inline(para, txt)

            tcPr = cell._tc.get_or_add_tcPr()
            # 清除可能重复的 tcMar（三线表不使用底色 shd）
            for old in tcPr.findall(qn("w:tcMar")):
                tcPr.remove(old)
            tcMar = _ET.SubElement(tcPr, qn("w:tcMar"))
            for side in ("top", "left", "bottom", "right"):
                m = _ET.SubElement(tcMar, qn(f"w:{side}"))
                m.set(qn("w:w"), "80"); m.set(qn("w:type"), "dxa")

            if is_hdr:
                _set_cell_header_border(cell)
                for run in para.runs:
                    run.bold = True   # 表头加粗，黑字无底色（三线表标准样式）

# ── 5e. 主转换函数 ──
_DISP_SINGLE_RE = re.compile(r'^\s*\$\$([\s\S]+?)\$\$\s*$')

def md_to_docx(md_text: str, out_path: str) -> None:
    print("📝 解析 Markdown 并构建 Word 文档…")

    # 检测是否有 Office XSL（影响公式渲染质量）
    xsl_ok = any(os.path.exists(p) for p in [
        "/Applications/Microsoft Word.app/Contents/Resources/MML2OMML.XSL",
        os.path.expanduser("~/Library/Application Support/Microsoft/Office/MML2OMML.XSL"),
        r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL",
        r"C:\Program Files (x86)\Microsoft Office\root\Office16\MML2OMML.XSL",
    ])
    if xsl_ok:
        print("  ✅ 检测到 Office for Mac，公式将转为原生 OMML（可在 Word 中编辑）。")
    else:
        print("  ℹ️  未检测到 Office XSL，公式使用 MathML 内嵌（Word 2013+ 可显示）。")

    doc = Document()

    # A4 页面 + 页边距
    sec = doc.sections[0]
    sec.page_height = Cm(29.7); sec.page_width = Cm(21.0)
    sec.top_margin = sec.bottom_margin = Cm(2.5)
    sec.left_margin = sec.right_margin = Cm(3.0)

    _set_doc_defaults(doc)

    lines    = md_text.splitlines()
    i        = 0
    in_code  = False
    code_buf: list[str] = []

    while i < len(lines):
        line = lines[i]

        # ── 围栏代码块 ────────────────────────────────────────
        if line.strip().startswith("```"):
            if not in_code:
                in_code = True; code_buf = []
            else:
                in_code = False
                _add_code_block(doc, "\n".join(code_buf))
            i += 1; continue
        if in_code:
            code_buf.append(line); i += 1; continue

        # ── 多行显示公式（$$ 单独占一行） ──────────────────────
        if line.strip() == "$$":
            buf = []
            i += 1
            while i < len(lines) and lines[i].strip() != "$$":
                buf.append(lines[i]); i += 1
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _append_formula(para, "$$" + "\n".join(buf) + "$$")
            i += 1; continue

        # ── 单行显示公式 $$expr$$ ───────────────────────────────
        if _DISP_SINGLE_RE.match(line):
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _append_formula(para, line.strip())
            i += 1; continue

        # ── 标题 #...###### ────────────────────────────────────
        hm = re.match(r'^(#{1,6})\s+(.*)', line)
        if hm:
            level = len(hm.group(1))
            try:   para = doc.add_paragraph(style=f"Heading {min(level, 6)}")
            except KeyError: para = doc.add_paragraph()
            _add_inline(para, hm.group(2).strip())
            i += 1; continue

        # ── 水平线 ─────────────────────────────────────────────
        if re.match(r'^[-*_]{3,}\s*$', line.strip()):
            para = doc.add_paragraph()
            pPr  = para._p.get_or_add_pPr()
            pBdr = _ET.SubElement(pPr, qn("w:pBdr"))
            bot  = _ET.SubElement(pBdr, qn("w:bottom"))
            bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"),    "6")
            bot.set(qn("w:space"), "1");    bot.set(qn("w:color"), "AAAAAA")
            i += 1; continue

        # ── 引用 > ─────────────────────────────────────────────
        if line.startswith(">"):
            try:   para = doc.add_paragraph(style="Quote")
            except KeyError: para = doc.add_paragraph()
            _add_inline(para, re.sub(r'^>\s?', '', line))
            i += 1; continue

        # ── 无序列表 - / * / + ─────────────────────────────────
        ulm = re.match(r'^(\s*)[-*+]\s+(.*)', line)
        if ulm:
            depth = len(ulm.group(1)) // 2
            try:   para = doc.add_paragraph(style="List Bullet")
            except KeyError: para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(0.75 * (depth + 1))
            _add_inline(para, ulm.group(2)); i += 1; continue

        # ── 有序列表 1. 2. ─────────────────────────────────────
        olm = re.match(r'^(\s*)\d+\.\s+(.*)', line)
        if olm:
            depth = len(olm.group(1)) // 2
            try:   para = doc.add_paragraph(style="List Number")
            except KeyError: para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(0.75 * (depth + 1))
            _add_inline(para, olm.group(2)); i += 1; continue

        # ── 表格 | ─────────────────────────────────────────────
        if '|' in line and line.strip().startswith('|'):
            raw_rows = []
            while i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
                raw_rows.append(_parse_table_row(lines[i]))
                i += 1
            header, body, aligns = None, raw_rows, None
            if len(raw_rows) >= 2 and _is_sep_row(raw_rows[1]):
                header = raw_rows[0]
                aligns = [_col_align(c) for c in raw_rows[1]]
                body   = raw_rows[2:]
            _add_table_block(doc, header, body, aligns)
            continue

        # ── 空行 ───────────────────────────────────────────────
        if not line.strip():
            i += 1; continue

        # ── 普通段落 ───────────────────────────────────────────
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(6)
        _add_inline(para, line)
        i += 1

    doc.save(out_path)

    # XML 完整性校验（确保 Word 可以打开）
    try:
        with open(out_path, "rb") as fh:
            xml = zipfile.ZipFile(io.BytesIO(fh.read())).read("word/document.xml")
            _ET.fromstring(xml)
        print("  ✅ XML 结构校验通过，Word 可正常打开。")
    except Exception as e:
        print(f"  ⚠️  XML 校验警告：{e}")


# ─────────────────────────────────────────────────
# § 6  主流程
# ─────────────────────────────────────────────────
def main() -> None:
    print("=" * 55)
    print("  📄 Markdown → PDF / Word  v2.0")
    print("     中文 · LaTeX · 有机化学式")
    print("=" * 55)

    md_text  = read_clipboard()
    md_text  = maybe_strip_ai_preamble(md_text)
    fmt      = ask_format()
    theme    = ask_pdf_theme() if fmt == "pdf" else "pingfang"
    ext      = ".pdf" if fmt == "pdf" else ".docx"
    name     = ask_filename(ext)

    desktop  = os.path.expanduser("~/Desktop")
    os.makedirs(desktop, exist_ok=True)
    out_path = os.path.join(desktop, name + ext)
    print()

    if fmt == "pdf":
        label = ("苹方 + Helvetica Neue" if theme == "pingfang"
                 else "Noto Serif CJK SC + Times New Roman")
        print(f"🎨 字体主题：{label}")
        print("🔄 解析 Markdown…")
        html = build_full_html(md_to_html_body(md_text), theme=theme)
        html_to_pdf(html, out_path)
    else:
        print("🔤 字体：宋体 + Times New Roman")
        md_to_docx(md_text, out_path)

    print(f"\n✅ 已保存至桌面：{out_path}")
    try:
        if sys.platform == "darwin":
            subprocess.run(["open", "-R", out_path], check=False)
        elif sys.platform == "win32":
            subprocess.run(["explorer", f"/select,{out_path}"], check=False)
    except Exception:
        pass


if __name__ == "__main__":
    main()
